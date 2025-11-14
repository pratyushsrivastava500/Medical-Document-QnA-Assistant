"""
Generate comprehensive project report in DOCX format
"""

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Installing python-docx...")
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement


def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()


def create_project_report():
    """Create comprehensive project report"""
    doc = Document()
    
    # Title Page
    title = doc.add_heading('Medical Document Analysis Assistant', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('AI-Powered RAG System for Healthcare', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run('Complete Project Report')
    run.bold = True
    run.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Deployment Platform: Streamlit.io')
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_page_break(doc)
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. Project Overview',
        '3. System Architecture',
        '4. Technology Stack',
        '5. Core Features',
        '6. System Components',
        '7. Workflow Description',
        '8. Deployment Platform',
        '9. Installation & Setup',
        '10. Usage Guide',
        '11. Configuration Details',
        '12. Future Enhancements'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    add_page_break(doc)
    
    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'The Medical Document Analysis Assistant is an intelligent AI-powered system designed for '
        'healthcare organizations to analyze medical documents, extract insights, and generate '
        'professional reports. Built using Retrieval-Augmented Generation (RAG) technology, the '
        'system ensures accurate, grounded responses with zero hallucination by grounding all '
        'answers in the actual document content.'
    )
    
    doc.add_paragraph('Key Capabilities:')
    capabilities = [
        'Multi-format document processing (PDF, Word, Excel, Images)',
        'Interactive Q&A with source citations',
        'Professional medical report generation with 5 comprehensive sections',
        'Google Drive integration for seamless document management',
        'PDF export functionality for reports',
        'Vector-based semantic search using FAISS',
        'Context-aware responses using advanced LLM (GPT-4.1-Nano)'
    ]
    for cap in capabilities:
        p = doc.add_paragraph(cap, style='List Bullet')
    
    add_page_break(doc)
    
    # 2. Project Overview
    doc.add_heading('2. Project Overview', level=1)
    
    doc.add_heading('Problem Statement', level=2)
    doc.add_paragraph(
        'Healthcare professionals often need to analyze large volumes of medical documents to extract '
        'specific information, generate reports, and make informed decisions. Traditional manual '
        'document review is time-consuming and prone to human error. This system automates the process '
        'using AI while maintaining accuracy through source citations.'
    )
    
    doc.add_heading('Solution Approach', level=2)
    doc.add_paragraph(
        'The system implements a Retrieval-Augmented Generation (RAG) pipeline that:'
    )
    solutions = [
        'Processes and indexes medical documents using advanced NLP techniques',
        'Converts documents into semantic embeddings for efficient retrieval',
        'Stores embeddings in a vector database (FAISS) for fast similarity search',
        'Retrieves relevant context based on user queries',
        'Generates accurate responses using Euriai LLM API (GPT-4.1-Nano)',
        'Maintains conversation history for context-aware interactions',
        'Provides source citations for transparency and verification'
    ]
    for sol in solutions:
        doc.add_paragraph(sol, style='List Bullet')
    
    add_page_break(doc)
    
    # 3. System Architecture
    doc.add_heading('3. System Architecture', level=1)
    
    doc.add_paragraph(
        'The system follows a modular architecture with clear separation of concerns:'
    )
    
    doc.add_heading('Architecture Diagram', level=2)
    doc.add_paragraph(
        'The following diagram illustrates the system components and data flow:'
    )
    
    # ASCII Architecture Diagram
    arch_text = """
    ┌─────────────────────────────────────────────────────────────────┐
    │                        USER INTERFACE                            │
    │                      (Streamlit App)                             │
    │  ┌─────────────────┐                    ┌────────────────────┐ │
    │  │  Q&A Mode       │                    │  Report Generation │ │
    │  │  - Ask questions│                    │  - 5 Sections      │ │
    │  │  - Get citations│                    │  - PDF Export      │ │
    │  └─────────────────┘                    └────────────────────┘ │
    └───────────────────────────┬──────────────────────────────────────┘
                                │
                                ▼
    ┌─────────────────────────────────────────────────────────────────┐
    │                    DOCUMENT PROCESSING LAYER                     │
    │  ┌──────────────┐  ┌──────────────┐  ┌─────────────────────┐  │
    │  │ PDF Parser   │  │ DOCX Parser  │  │ Image OCR (Pytesseract)│
    │  │ (PyPDF2)     │  │ (python-docx)│  │                     │  │
    │  └──────────────┘  └──────────────┘  └─────────────────────┘  │
    │  ┌──────────────────────────────────────────────────────────┐  │
    │  │         Text Chunking & Preprocessing                     │  │
    │  │         (Semantic Chunking with Overlap)                  │  │
    │  └──────────────────────────────────────────────────────────┘  │
    └───────────────────────────┬──────────────────────────────────────┘
                                │
                                ▼
    ┌─────────────────────────────────────────────────────────────────┐
    │                    EMBEDDING & STORAGE LAYER                     │
    │  ┌──────────────────────────────────────────────────────────┐  │
    │  │     Sentence Transformers (all-MiniLM-L6-v2)              │  │
    │  │     Converts text chunks to 384-dimensional vectors       │  │
    │  └──────────────────────────────────────────────────────────┘  │
    │                              │                                   │
    │                              ▼                                   │
    │  ┌──────────────────────────────────────────────────────────┐  │
    │  │           FAISS Vector Database                           │  │
    │  │           (Facebook AI Similarity Search)                 │  │
    │  │           - Fast similarity search                        │  │
    │  │           - Efficient retrieval (top-k)                   │  │
    │  └──────────────────────────────────────────────────────────┘  │
    └───────────────────────────┬──────────────────────────────────────┘
                                │
                                ▼
    ┌─────────────────────────────────────────────────────────────────┐
    │                     RAG PIPELINE LAYER                           │
    │  ┌──────────────────────────────────────────────────────────┐  │
    │  │  1. Query Processing & Embedding                          │  │
    │  └──────────────────────────────────────────────────────────┘  │
    │                              │                                   │
    │                              ▼                                   │
    │  ┌──────────────────────────────────────────────────────────┐  │
    │  │  2. Semantic Search in Vector DB                          │  │
    │  │     (Retrieve top 5 relevant chunks)                      │  │
    │  └──────────────────────────────────────────────────────────┘  │
    │                              │                                   │
    │                              ▼                                   │
    │  ┌──────────────────────────────────────────────────────────┐  │
    │  │  3. Context Augmentation                                  │  │
    │  │     (Combine query + retrieved context + conversation)    │  │
    │  └──────────────────────────────────────────────────────────┘  │
    │                              │                                   │
    │                              ▼                                   │
    │  ┌──────────────────────────────────────────────────────────┐  │
    │  │  4. LLM Generation (Euriai API - GPT-4.1-Nano)            │  │
    │  └──────────────────────────────────────────────────────────┘  │
    └───────────────────────────┬──────────────────────────────────────┘
                                │
                                ▼
    ┌─────────────────────────────────────────────────────────────────┐
    │                    EXTERNAL INTEGRATIONS                         │
    │  ┌──────────────────┐              ┌─────────────────────────┐ │
    │  │ Google Drive API │              │  PDF Export (ReportLab) │ │
    │  │ - File upload    │              │  - Professional reports │ │
    │  │ - File download  │              │  - Styled documents     │ │
    │  └──────────────────┘              └─────────────────────────┘ │
    └─────────────────────────────────────────────────────────────────┘
    """
    
    p = doc.add_paragraph(arch_text)
    p.style = 'Normal'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
    
    add_page_break(doc)
    
    # 4. Technology Stack
    doc.add_heading('4. Technology Stack', level=1)
    
    doc.add_heading('4.1 Programming Languages', level=2)
    doc.add_paragraph('Python 3.8+: Core programming language for the entire application')
    
    doc.add_heading('4.2 Web Framework', level=2)
    doc.add_paragraph('Streamlit 1.29.0: Interactive web application framework for building the user interface')
    
    doc.add_heading('4.3 AI & Machine Learning', level=2)
    ml_stack = [
        ('Euriai API (GPT-4.1-Nano)', 'Advanced language model for natural language understanding and generation'),
        ('Sentence Transformers 2.3.1', 'Pre-trained model (all-MiniLM-L6-v2) for generating semantic embeddings'),
        ('FAISS (CPU) 1.7.4', 'Facebook AI Similarity Search for efficient vector similarity search'),
        ('PyTorch 2.1.2', 'Deep learning framework (dependency for Sentence Transformers)')
    ]
    for tech, desc in ml_stack:
        p = doc.add_paragraph()
        p.add_run(tech + ': ').bold = True
        p.add_run(desc)
    
    doc.add_heading('4.4 Document Processing', level=2)
    doc_stack = [
        ('PyPDF2 3.0.1', 'PDF document parsing and text extraction'),
        ('python-docx 1.1.0', 'Microsoft Word document processing'),
        ('openpyxl 3.1.2', 'Excel spreadsheet reading and processing'),
        ('Pillow (PIL) 10.1.0', 'Image processing and manipulation'),
        ('pytesseract 0.3.10', 'OCR (Optical Character Recognition) for extracting text from images'),
        ('ReportLab 4.0.7', 'Professional PDF report generation with styling')
    ]
    for tech, desc in doc_stack:
        p = doc.add_paragraph()
        p.add_run(tech + ': ').bold = True
        p.add_run(desc)
    
    doc.add_heading('4.5 Cloud Storage & APIs', level=2)
    cloud_stack = [
        ('Google Drive API', 'Cloud storage integration for file upload and management'),
        ('google-auth 2.25.2', 'Authentication for Google services'),
        ('google-auth-oauthlib 1.2.0', 'OAuth 2.0 authentication flow'),
        ('google-api-python-client 2.111.0', 'Google API client libraries')
    ]
    for tech, desc in cloud_stack:
        p = doc.add_paragraph()
        p.add_run(tech + ': ').bold = True
        p.add_run(desc)
    
    doc.add_heading('4.6 Utilities & Supporting Libraries', level=2)
    util_stack = [
        ('python-dotenv 1.0.0', 'Environment variable management for secure credential storage'),
        ('requests 2.31.0', 'HTTP library for API communications'),
        ('numpy 1.24.3', 'Numerical computing for vector operations')
    ]
    for tech, desc in util_stack:
        p = doc.add_paragraph()
        p.add_run(tech + ': ').bold = True
        p.add_run(desc)
    
    doc.add_heading('4.7 Deployment Platform', level=2)
    p = doc.add_paragraph()
    p.add_run('Streamlit.io: ').bold = True
    p.add_run(
        'Cloud-based deployment platform specifically designed for Streamlit applications. '
        'Provides seamless deployment directly from GitHub repositories with automatic builds, '
        'HTTPS security, custom domain support, and built-in scaling capabilities.'
    )
    
    deployment_features = [
        'One-click deployment from GitHub',
        'Automatic rebuilds on code changes',
        'Built-in SSL/TLS certificates',
        'Environment variable management',
        'Resource monitoring and analytics',
        'Custom subdomain (*.streamlit.app)',
        'Free tier available for public projects'
    ]
    doc.add_paragraph('Key Features:')
    for feature in deployment_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    add_page_break(doc)
    
    # 5. Core Features
    doc.add_heading('5. Core Features', level=1)
    
    doc.add_heading('5.1 Document Upload & Processing', level=2)
    doc.add_paragraph(
        'Multi-source document ingestion supporting:'
    )
    upload_features = [
        'Local file upload (PDF, DOCX, XLSX, images)',
        'Google Drive integration with OAuth 2.0 authentication',
        'Batch processing of multiple documents',
        'Automatic text extraction from various formats',
        'OCR processing for scanned documents and images',
        'Persistent vector database storage'
    ]
    for feature in upload_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('5.2 Interactive Q&A Mode', level=2)
    doc.add_paragraph(
        'Conversational interface with advanced capabilities:'
    )
    qa_features = [
        'Natural language question answering',
        'Context-aware responses using conversation history',
        'Source citations with document references and page numbers',
        'Confidence scoring for answers',
        'Multi-turn conversations with memory',
        'Error handling with graceful degradation (citation suppression on API errors)'
    ]
    for feature in qa_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('5.3 Medical Report Generation', level=2)
    doc.add_paragraph(
        'Comprehensive report creation with structured sections:'
    )
    report_sections = [
        'Introduction: Patient overview and medical context',
        'Clinical Findings: Detailed analysis of symptoms, test results, and observations',
        'Diagnosis: Medical conditions identified with supporting evidence',
        'Treatment Plan: Recommended interventions, medications, and procedures',
        'Summary: Concise overview and key takeaways'
    ]
    doc.add_paragraph('Report Sections:')
    for section in report_sections:
        doc.add_paragraph(section, style='List Bullet')
    
    doc.add_paragraph('Report Features:')
    report_features = [
        'Professional formatting with consistent styling',
        'Automatic duplicate heading removal',
        'Source citations for each section',
        'PDF export with ReportLab integration',
        'Downloadable reports with metadata'
    ]
    for feature in report_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    add_page_break(doc)
    
    # 6. System Components
    doc.add_heading('6. System Components', level=1)
    
    components = [
        ('app.py (~1448 lines)', 
         'Main Streamlit application handling UI, user interactions, document upload, '
         'mode selection (Q&A/Report), conversation display, report generation interface, '
         'PDF download functionality, and error handling.'),
        
        ('services/rag_pipeline.py',
         'Core RAG implementation managing query processing, context retrieval from vector store, '
         'prompt construction, LLM API calls, and response parsing with citation extraction.'),
        
        ('services/llm_client.py',
         'Euriai API client handling authentication, request formatting, response parsing, '
         'error handling (including 403 error detection), and retry logic.'),
        
        ('services/conversation_memory.py',
         'Conversation history management implementing sliding window context, message storage, '
         'context size optimization, and conversation state persistence.'),
        
        ('services/report_generator.py (~881 lines)',
         'Medical report generation engine with section-specific prompt templates, LLM calls for '
         'each section, duplicate heading removal, text cleaning (hyphenation, footer removal), '
         'and citation extraction.'),
        
        ('services/google_drive.py',
         'Google Drive API integration handling OAuth 2.0 authentication, file listing, '
         'file upload/download, metadata management, and caching.'),
        
        ('utils/document_processor.py',
         'Multi-format document processing with PDF parsing (PyPDF2), DOCX parsing (python-docx), '
         'Excel parsing (openpyxl), image OCR (pytesseract), and text extraction.'),
        
        ('utils/embeddings.py',
         'Text embedding generation using Sentence Transformers model (all-MiniLM-L6-v2), '
         'batch processing, and vector normalization.'),
        
        ('utils/vector_store.py',
         'FAISS vector database management handling index creation, document addition, '
         'similarity search (top-k retrieval), index persistence (save/load), and metadata storage.'),
        
        ('config/config.py',
         'Configuration management loading environment variables (.env), API credentials, '
         'model settings, and file paths.')
    ]
    
    for comp_name, comp_desc in components:
        p = doc.add_paragraph()
        p.add_run(comp_name + ': ').bold = True
        p.add_run(comp_desc)
    
    add_page_break(doc)
    
    # 7. Workflow Description
    doc.add_heading('7. Workflow Description', level=1)
    
    doc.add_heading('7.1 Document Upload & Indexing Workflow', level=2)
    workflow_steps = [
        ('Step 1: Document Upload', 
         'User uploads documents via local file upload or Google Drive integration'),
        
        ('Step 2: Format Detection', 
         'System automatically detects file format (PDF, DOCX, XLSX, image)'),
        
        ('Step 3: Text Extraction', 
         'Appropriate parser extracts text content (OCR for images)'),
        
        ('Step 4: Text Chunking', 
         'Text is split into semantic chunks with overlap for context preservation'),
        
        ('Step 5: Embedding Generation', 
         'Sentence Transformers converts each chunk to 384-dimensional vector'),
        
        ('Step 6: Vector Storage', 
         'FAISS stores vectors with metadata (document name, page number, content)'),
        
        ('Step 7: Index Persistence', 
         'Vector database is saved to disk for future sessions')
    ]
    
    for step_name, step_desc in workflow_steps:
        p = doc.add_paragraph()
        p.add_run(step_name + ': ').bold = True
        p.add_run(step_desc)
    
    doc.add_heading('7.2 Q&A Mode Workflow', level=2)
    qa_workflow = [
        ('Step 1: User Query', 
         'User enters natural language question in chat interface'),
        
        ('Step 2: Query Embedding', 
         'Question is converted to embedding vector using same model'),
        
        ('Step 3: Similarity Search', 
         'FAISS performs cosine similarity search to find top 5 relevant chunks'),
        
        ('Step 4: Context Aggregation', 
         'Retrieved chunks are combined with conversation history'),
        
        ('Step 5: Prompt Construction', 
         'System prompt + user query + context + conversation history'),
        
        ('Step 6: LLM Generation', 
         'Euriai API (GPT-4.1-Nano) generates response based on context'),
        
        ('Step 7: Citation Extraction', 
         'System extracts source documents and page numbers from response'),
        
        ('Step 8: Response Display', 
         'Answer is shown with citations (suppressed if API error detected)'),
        
        ('Step 9: Memory Update', 
         'Conversation history is updated with Q&A pair')
    ]
    
    for step_name, step_desc in qa_workflow:
        p = doc.add_paragraph()
        p.add_run(step_name + ': ').bold = True
        p.add_run(step_desc)
    
    doc.add_heading('7.3 Report Generation Workflow', level=2)
    report_workflow = [
        ('Step 1: Mode Selection', 
         'User selects Report Generation mode from sidebar'),
        
        ('Step 2: Patient Info (Optional)', 
         'User can provide patient name and ID for report header'),
        
        ('Step 3: Section Generation', 
         'System generates each of 5 sections sequentially using LLM'),
        
        ('Step 4: Content Processing', 
         'Each section undergoes duplicate heading removal and text cleaning'),
        
        ('Step 5: Citation Extraction', 
         'Sources are extracted for each section (suppressed on errors)'),
        
        ('Step 6: Report Display', 
         'All sections are displayed in formatted markdown'),
        
        ('Step 7: PDF Export', 
         'User can download report as PDF using ReportLab'),
        
        ('Step 8: Report Storage', 
         'Report is temporarily stored in session state for download')
    ]
    
    for step_name, step_desc in report_workflow:
        p = doc.add_paragraph()
        p.add_run(step_name + ': ').bold = True
        p.add_run(step_desc)
    
    add_page_break(doc)
    
    # 8. Deployment Platform Details
    doc.add_heading('8. Deployment Platform: Streamlit.io', level=1)
    
    doc.add_heading('8.1 Why Streamlit.io?', level=2)
    doc.add_paragraph(
        'Streamlit.io is the official cloud platform for deploying Streamlit applications, '
        'offering seamless integration and optimized performance for Python-based data apps.'
    )
    
    advantages = [
        'Zero Configuration: No Docker, Kubernetes, or complex DevOps setup required',
        'GitHub Integration: Automatic deployment from repository with CI/CD',
        'Instant Updates: Changes pushed to GitHub are automatically deployed',
        'Secure by Default: Built-in HTTPS, environment variable encryption',
        'Cost-Effective: Free tier for public projects with generous resource limits',
        'Streamlit-Optimized: Platform designed specifically for Streamlit apps',
        'Monitoring Tools: Built-in analytics and resource usage tracking',
        'Community Support: Large user community and extensive documentation'
    ]
    
    doc.add_paragraph('Advantages:')
    for adv in advantages:
        doc.add_paragraph(adv, style='List Bullet')
    
    doc.add_heading('8.2 Deployment Configuration', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Repository Structure for Deployment:').bold = True
    
    deploy_files = [
        ('requirements.txt', 'Python dependencies (automatically installed)'),
        ('.streamlit/config.toml', 'Streamlit configuration (theme, server settings)'),
        ('.env (Secrets)', 'Environment variables (configured in Streamlit Cloud dashboard)'),
        ('app.py', 'Main entry point (auto-detected by Streamlit)')
    ]
    
    for file_name, file_desc in deploy_files:
        p = doc.add_paragraph()
        p.add_run(file_name + ': ').bold = True
        p.add_run(file_desc)
    
    doc.add_heading('8.3 Resource Requirements', level=2)
    
    resources = [
        ('Memory', '2GB RAM minimum (for FAISS vector operations)'),
        ('Storage', '1GB for vector database and cached files'),
        ('CPU', '2 cores recommended for concurrent user handling'),
        ('Python Version', '3.8+ (specified in runtime.txt)')
    ]
    
    for res_name, res_value in resources:
        p = doc.add_paragraph()
        p.add_run(res_name + ': ').bold = True
        p.add_run(res_value)
    
    add_page_break(doc)
    
    # 9. Installation & Setup
    doc.add_heading('9. Installation & Setup', level=1)
    
    doc.add_heading('9.1 Local Development Setup', level=2)
    
    setup_steps = [
        'Clone the repository',
        'Create Python virtual environment: python -m venv venv',
        'Activate environment: venv\\Scripts\\activate (Windows) or source venv/bin/activate (Linux/Mac)',
        'Install dependencies: pip install -r requirements.txt',
        'Install Tesseract OCR: Download from https://github.com/tesseract-ocr/tesseract',
        'Create .env file with API credentials (EURIAI_API_KEY, GOOGLE_DRIVE_CREDENTIALS)',
        'Run application: streamlit run app.py',
        'Access at http://localhost:8501'
    ]
    
    for i, step in enumerate(setup_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('9.2 Streamlit.io Deployment', level=2)
    
    deploy_steps = [
        'Push code to GitHub repository',
        'Sign up/Login to https://streamlit.io',
        'Click "New app" from dashboard',
        'Connect GitHub account and select repository',
        'Specify branch (main) and entry point (app.py)',
        'Add secrets in Advanced Settings (EURIAI_API_KEY, etc.)',
        'Click "Deploy" button',
        'Wait for build completion (~5 minutes)',
        'Access app at provided URL (*.streamlit.app)'
    ]
    
    for i, step in enumerate(deploy_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    add_page_break(doc)
    
    # 10. Usage Guide
    doc.add_heading('10. Usage Guide', level=1)
    
    doc.add_heading('10.1 Document Upload', level=2)
    doc.add_paragraph(
        'Upload medical documents through local file browser or Google Drive integration. '
        'Supported formats include PDF, DOCX, XLSX, PNG, JPG. The system automatically processes '
        'and indexes documents for querying.'
    )
    
    doc.add_heading('10.2 Q&A Mode', level=2)
    doc.add_paragraph(
        'Ask questions in natural language about uploaded documents. The system retrieves relevant '
        'context and generates accurate answers with source citations. Conversation history is '
        'maintained for follow-up questions.'
    )
    
    doc.add_heading('10.3 Report Generation', level=2)
    doc.add_paragraph(
        'Generate comprehensive medical reports covering Introduction, Clinical Findings, Diagnosis, '
        'Treatment Plan, and Summary. Optionally provide patient information. Export reports as PDF '
        'for record-keeping or sharing.'
    )
    
    add_page_break(doc)
    
    # 11. Configuration Details
    doc.add_heading('11. Configuration Details', level=1)
    
    doc.add_heading('11.1 Environment Variables', level=2)
    env_vars = [
        ('EURIAI_API_KEY', 'API key for Euriai LLM service (required)'),
        ('GOOGLE_DRIVE_CREDENTIALS', 'OAuth 2.0 credentials JSON for Google Drive (optional)'),
        ('FAISS_INDEX_PATH', 'Path to vector database storage (default: data/vector_db/)'),
        ('CACHE_DIR', 'Directory for cached files (default: data/gdrive_cache/)')
    ]
    
    for var_name, var_desc in env_vars:
        p = doc.add_paragraph()
        p.add_run(var_name + ': ').bold = True
        p.add_run(var_desc)
    
    doc.add_heading('11.2 Model Configuration', level=2)
    model_config = [
        ('LLM Model', 'gpt-4.1-nano (Euriai API)'),
        ('Embedding Model', 'sentence-transformers/all-MiniLM-L6-v2'),
        ('Vector Dimension', '384 (from embedding model)'),
        ('Top-K Retrieval', '5 chunks per query'),
        ('Max Context Length', '4096 tokens'),
        ('Temperature', '0.7 (balanced creativity and consistency)')
    ]
    
    for config_name, config_value in model_config:
        p = doc.add_paragraph()
        p.add_run(config_name + ': ').bold = True
        p.add_run(config_value)
    
    add_page_break(doc)
    
    # 12. Future Enhancements
    doc.add_heading('12. Future Enhancements', level=1)
    
    enhancements = [
        'Multi-language support for international healthcare organizations',
        'Advanced analytics dashboard with usage statistics and trends',
        'Integration with Electronic Health Record (EHR) systems',
        'Real-time collaboration features for team-based report generation',
        'Mobile-responsive design for tablet and smartphone access',
        'Voice-to-text input for hands-free operation',
        'Custom report templates for different medical specialties',
        'Automated medical coding (ICD-10, CPT) integration',
        'HIPAA compliance certification and audit logging',
        'Database integration for persistent report storage',
        'User role management and access control',
        'Advanced visualization options for medical data',
        'Integration with medical imaging systems (DICOM)',
        'Multilingual report generation',
        'Automated quality checks and medical terminology validation'
    ]
    
    for enh in enhancements:
        doc.add_paragraph(enh, style='List Bullet')
    
    add_page_break(doc)
    
    # Conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'The Medical Document Analysis Assistant represents a significant advancement in healthcare '
        'document processing and analysis. By combining state-of-the-art AI technologies including '
        'RAG architecture, advanced language models, and semantic search, the system provides '
        'healthcare professionals with a powerful tool for efficient document analysis and report '
        'generation.'
    )
    
    doc.add_paragraph(
        'The deployment on Streamlit.io ensures accessibility, scalability, and ease of maintenance, '
        'making it suitable for both small clinics and large healthcare organizations. The modular '
        'architecture and comprehensive technology stack provide a solid foundation for future '
        'enhancements and customization to meet specific organizational needs.'
    )
    
    doc.add_paragraph(
        'With its focus on accuracy through source citations, user-friendly interface, and '
        'professional report generation capabilities, this system stands ready to transform '
        'medical document workflows and improve healthcare delivery efficiency.'
    )
    
    # Save document
    output_path = 'd:/Hiring/Medical_Document_Analysis_Project_Report.docx'
    doc.save(output_path)
    print(f'✓ Project report created successfully: {output_path}')
    return output_path


if __name__ == '__main__':
    create_project_report()
