"""
Document Processing Module
Handles parsing and processing of various document formats
"""
import os
import io
from typing import List, Dict, Any
from pathlib import Path

# Document processing libraries - all optional
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    try:
        import pypdf as PyPDF2
        PDF_AVAILABLE = True
    except ImportError:
        print("Warning: PyPDF2/pypdf not installed. PDF support disabled.")
        PyPDF2 = None
        PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    print("Warning: python-docx not installed. Word document support disabled.")
    Document = None
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    print("Warning: pandas/openpyxl not installed. Excel support disabled.")
    pd = None
    EXCEL_AVAILABLE = False

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    print("Warning: Pillow/pytesseract not installed. OCR support disabled.")
    Image = None
    pytesseract = None
    OCR_AVAILABLE = False

from langchain.text_splitter import RecursiveCharacterTextSplitter
from config.config import Config


class DocumentProcessor:
    """Handles document parsing, extraction, and chunking"""
    
    def __init__(self):
        self.chunk_size = Config.CHUNK_SIZE
        self.chunk_overlap = Config.CHUNK_OVERLAP
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def process_document(self, file_path: str, file_name: str = None, 
                        source_type: str = "uploaded") -> List[Dict[str, Any]]:
        """
        Process a document and return chunked text with metadata
        
        Args:
            file_path: Path to the document
            file_name: Original filename
            source_type: "uploaded" or "gdrive"
            
        Returns:
            List of document chunks with metadata
        """
        if file_name is None:
            file_name = os.path.basename(file_path)
        
        # Extract text based on file type
        text = self._extract_text(file_path, file_name)
        
        if not text or len(text.strip()) == 0:
            return []
        
        # Split into chunks
        chunks = self.text_splitter.split_text(text)
        
        # Create documents with metadata
        documents = []
        for i, chunk in enumerate(chunks):
            doc = {
                "content": chunk,
                "metadata": {
                    "source": file_name,
                    "chunk_id": i,
                    "total_chunks": len(chunks),
                    "source_type": source_type,
                    "file_path": file_path
                }
            }
            documents.append(doc)
        
        return documents
    
    def _extract_text(self, file_path: str, file_name: str) -> str:
        """Extract text from various file formats"""
        ext = Path(file_name).suffix.lower()
        
        try:
            if ext == '.pdf':
                return self._extract_from_pdf(file_path)
            elif ext in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            elif ext in ['.xlsx', '.xls']:
                return self._extract_from_excel(file_path)
            elif ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                return self._extract_from_image(file_path)
            elif ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                print(f"Unsupported file format: {ext}")
                return ""
        except Exception as e:
            print(f"Error extracting text from {file_name}: {str(e)}")
            return ""
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""
        if not PDF_AVAILABLE:
            return "Error: PDF support not available. Install pypdf: pip install pypdf"
        
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}"
        except Exception as e:
            print(f"Error reading PDF: {str(e)}")
        return text
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from Word document"""
        if not DOCX_AVAILABLE:
            return "Error: Word document support not available. Install python-docx: pip install python-docx"
        
        try:
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text += f"\n{row_text}"
            
            return text
        except Exception as e:
            print(f"Error reading DOCX: {str(e)}")
            return ""
    
    def _extract_from_excel(self, file_path: str) -> str:
        """Extract text from Excel file"""
        if not EXCEL_AVAILABLE:
            return "Error: Excel support not available. Install pandas and openpyxl"
        
        try:
            excel_file = pd.ExcelFile(file_path)
            text = ""
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                text += f"\n--- Sheet: {sheet_name} ---\n"
                text += df.to_string(index=False)
                text += "\n"
            
            return text
        except Exception as e:
            print(f"Error reading Excel: {str(e)}")
            return ""
    
    def _extract_from_image(self, file_path: str) -> str:
        """Extract text from image using OCR"""
        if not OCR_AVAILABLE:
            return "Error: OCR support not available. Install Pillow and pytesseract"
        
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            print(f"Error reading image with OCR: {str(e)}")
            print("Make sure Tesseract is installed and in PATH")
            return ""
    
    def process_file_object(self, file_obj, file_name: str, 
                           source_type: str = "uploaded") -> List[Dict[str, Any]]:
        """
        Process a file object (e.g., from Streamlit upload)
        
        Args:
            file_obj: File-like object
            file_name: Name of the file
            source_type: Source type identifier
            
        Returns:
            List of document chunks with metadata
        """
        # Save to temporary location
        temp_path = os.path.join(Config.UPLOAD_DIR, file_name)
        
        with open(temp_path, 'wb') as f:
            f.write(file_obj.getbuffer())
        
        # Process the saved file
        return self.process_document(temp_path, file_name, source_type)
