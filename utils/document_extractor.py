"""
Document Extractor Module
Advanced extraction of tables, images, charts, and specific sections from documents
"""
import os
from typing import List, Dict, Any, Optional
from pathlib import Path
import io

# PDF extraction
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    try:
        import pypdf as PyPDF2
        PDF_AVAILABLE = True
    except ImportError:
        PyPDF2 = None
        PDF_AVAILABLE = False

# Word document extraction
try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    Document = None
    DOCX_AVAILABLE = False

# Excel extraction
try:
    import pandas as pd
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    pd = None
    EXCEL_AVAILABLE = False

# Image extraction from PDFs
try:
    from PIL import Image
    import pytesseract
    IMAGE_AVAILABLE = True
except ImportError:
    Image = None
    pytesseract = None
    IMAGE_AVAILABLE = False


class DocumentExtractor:
    """Advanced document content extraction"""
    
    def __init__(self):
        """Initialize document extractor"""
        pass
    
    def extract_tables_from_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract tables from PDF document
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            List of extracted tables with metadata
        """
        if not PDF_AVAILABLE:
            return [{"error": "PDF support not available"}]
        
        tables = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    
                    # Simple table detection: look for lines with multiple | or tab separators
                    lines = text.split('\n')
                    table_lines = []
                    in_table = False
                    
                    for line in lines:
                        # Detect table-like structure
                        if '|' in line or '\t' in line or '  ' in line:
                            table_lines.append(line)
                            in_table = True
                        elif in_table and line.strip() == '':
                            # End of table
                            if table_lines:
                                tables.append({
                                    "page": page_num + 1,
                                    "content": '\n'.join(table_lines),
                                    "type": "table",
                                    "rows": len(table_lines)
                                })
                            table_lines = []
                            in_table = False
                        elif in_table:
                            table_lines.append(line)
                    
                    # Catch remaining table
                    if table_lines:
                        tables.append({
                            "page": page_num + 1,
                            "content": '\n'.join(table_lines),
                            "type": "table",
                            "rows": len(table_lines)
                        })
        
        except Exception as e:
            print(f"Error extracting tables from PDF: {str(e)}")
        
        return tables
    
    def extract_tables_from_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract tables from Word document
        
        Args:
            file_path: Path to Word document
            
        Returns:
            List of extracted tables with metadata
        """
        if not DOCX_AVAILABLE:
            return [{"error": "Word document support not available"}]
        
        tables_data = []
        
        try:
            doc = Document(file_path)
            
            for idx, table in enumerate(doc.tables):
                # Extract table as formatted text
                table_text = []
                
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_text.append(' | '.join(row_data))
                
                tables_data.append({
                    "table_number": idx + 1,
                    "content": '\n'.join(table_text),
                    "type": "table",
                    "rows": len(table.rows),
                    "columns": len(table.rows[0].cells) if table.rows else 0
                })
        
        except Exception as e:
            print(f"Error extracting tables from Word: {str(e)}")
        
        return tables_data
    
    def extract_tables_from_excel(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract all sheets from Excel as tables
        
        Args:
            file_path: Path to Excel file
            
        Returns:
            List of sheets/tables with data
        """
        if not EXCEL_AVAILABLE:
            return [{"error": "Excel support not available"}]
        
        tables = []
        
        try:
            excel_file = pd.ExcelFile(file_path)
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Convert to formatted string
                table_content = df.to_string(index=False)
                
                tables.append({
                    "sheet_name": sheet_name,
                    "content": table_content,
                    "type": "table",
                    "rows": len(df),
                    "columns": len(df.columns),
                    "column_names": list(df.columns)
                })
        
        except Exception as e:
            print(f"Error extracting tables from Excel: {str(e)}")
        
        return tables
    
    def extract_all_tables(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract all tables from a document (auto-detect format)
        
        Args:
            file_path: Path to document
            
        Returns:
            List of all tables found
        """
        ext = Path(file_path).suffix.lower()
        
        if ext == '.pdf':
            return self.extract_tables_from_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return self.extract_tables_from_docx(file_path)
        elif ext in ['.xlsx', '.xls']:
            return self.extract_tables_from_excel(file_path)
        else:
            return []
    
    def extract_section_by_heading(self, file_path: str, 
                                   heading_keywords: List[str]) -> Dict[str, Any]:
        """
        Extract content under specific headings/sections
        
        Args:
            file_path: Path to document
            heading_keywords: Keywords to identify section headings
            
        Returns:
            Extracted section content
        """
        ext = Path(file_path).suffix.lower()
        
        if ext == '.pdf':
            return self._extract_section_from_pdf(file_path, heading_keywords)
        elif ext in ['.docx', '.doc']:
            return self._extract_section_from_docx(file_path, heading_keywords)
        else:
            return {"content": "", "found": False}
    
    def _extract_section_from_pdf(self, file_path: str, 
                                  heading_keywords: List[str]) -> Dict[str, Any]:
        """Extract section from PDF"""
        if not PDF_AVAILABLE:
            return {"error": "PDF support not available"}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                full_text = ""
                
                for page in pdf_reader.pages:
                    full_text += page.extract_text() + "\n"
                
                # Find sections
                lines = full_text.split('\n')
                section_content = []
                capturing = False
                
                for line in lines:
                    line_lower = line.lower()
                    
                    # Check if line contains heading keyword
                    if any(keyword.lower() in line_lower for keyword in heading_keywords):
                        capturing = True
                        section_content.append(line)
                        continue
                    
                    # Stop capturing at next major heading (all caps or ends with :)
                    if capturing:
                        if line.isupper() and len(line) > 3:
                            # Likely a new section
                            if not any(keyword.lower() in line_lower for keyword in heading_keywords):
                                break
                        section_content.append(line)
                
                return {
                    "content": '\n'.join(section_content),
                    "found": len(section_content) > 0
                }
        
        except Exception as e:
            print(f"Error extracting section from PDF: {str(e)}")
            return {"content": "", "found": False, "error": str(e)}
    
    def _extract_section_from_docx(self, file_path: str, 
                                   heading_keywords: List[str]) -> Dict[str, Any]:
        """Extract section from Word document"""
        if not DOCX_AVAILABLE:
            return {"error": "Word document support not available"}
        
        try:
            doc = Document(file_path)
            section_content = []
            capturing = False
            
            for para in doc.paragraphs:
                text = para.text
                text_lower = text.lower()
                
                # Check for heading
                if any(keyword.lower() in text_lower for keyword in heading_keywords):
                    capturing = True
                    section_content.append(text)
                    continue
                
                # Stop at next heading
                if capturing:
                    # Check if it's a style heading
                    if para.style.name.startswith('Heading'):
                        if not any(keyword.lower() in text_lower for keyword in heading_keywords):
                            break
                    section_content.append(text)
            
            return {
                "content": '\n'.join(section_content),
                "found": len(section_content) > 0
            }
        
        except Exception as e:
            print(f"Error extracting section from Word: {str(e)}")
            return {"content": "", "found": False, "error": str(e)}
    
    def extract_figures_and_charts_info(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract information about figures and charts from document
        
        Args:
            file_path: Path to document
            
        Returns:
            List of figure/chart references
        """
        ext = Path(file_path).suffix.lower()
        figures = []
        
        try:
            if ext == '.pdf':
                figures = self._extract_figure_refs_from_pdf(file_path)
            elif ext in ['.docx', '.doc']:
                figures = self._extract_figure_refs_from_docx(file_path)
        
        except Exception as e:
            print(f"Error extracting figures: {str(e)}")
        
        return figures
    
    def _extract_figure_refs_from_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract figure references from PDF"""
        if not PDF_AVAILABLE:
            return []
        
        figures = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    lines = text.split('\n')
                    
                    for line in lines:
                        line_lower = line.lower()
                        # Look for figure/chart/graph references
                        if any(keyword in line_lower for keyword in ['figure', 'fig.', 'chart', 'graph', 'diagram']):
                            figures.append({
                                "page": page_num + 1,
                                "reference": line.strip(),
                                "type": "figure_reference"
                            })
        
        except Exception as e:
            print(f"Error extracting figure references: {str(e)}")
        
        return figures
    
    def _extract_figure_refs_from_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract figure references from Word document"""
        if not DOCX_AVAILABLE:
            return []
        
        figures = []
        
        try:
            doc = Document(file_path)
            
            for para in doc.paragraphs:
                text_lower = para.text.lower()
                if any(keyword in text_lower for keyword in ['figure', 'fig.', 'chart', 'graph', 'diagram']):
                    figures.append({
                        "reference": para.text.strip(),
                        "type": "figure_reference"
                    })
        
        except Exception as e:
            print(f"Error extracting figure references from Word: {str(e)}")
        
        return figures
